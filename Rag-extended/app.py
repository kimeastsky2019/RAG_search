from fastapi import FastAPI, HTTPException, UploadFile, File, Depends, Form, status
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field
from contextlib import asynccontextmanager
from sqlmodel import select
from sqlalchemy import delete, func
from sqlmodel.ext.asyncio.session import AsyncSession
from datetime import timedelta
from typing import Optional
import uuid
import time
import os
import json
import httpx
from xai_sdk import AsyncClient

from config import XAI_API_KEY, XAI_MANAGEMENT_API_KEY, XAI_MODEL, COST_PER_1M_INPUT, COST_PER_1M_OUTPUT
from xai_sdk.proto import collections_pb2
from cache import cache_get, cache_set
from rag import run_rag
from database import init_db, get_session
from models import Collection, Document, User, UsageEvent
from ingest_folder import guess_content_type
from filters import build_metadata
from xai_helpers import extract_document_id, delete_collection_document
from auth_utils import verify_password, get_password_hash, create_access_token, ACCESS_TOKEN_EXPIRE_MINUTES

# Setup lifecycle management for DB init
@asynccontextmanager
async def lifespan(app: FastAPI):
    await init_db()
    
    # Create default user if not exists (for convenience)
    async for session in get_session():
        statement = select(User).where(User.email == "info@gngmeta.com")
        results = await session.exec(statement)
        user = results.first()
        if not user:
            default_user = User(
                email="info@gngmeta.com",
                hashed_password=get_password_hash("admin1234"), 
                full_name="GnG Admin"
            )
            session.add(default_user)
            await session.commit()
        break
        
    yield

app = FastAPI(title="Grok RAG Extended API", lifespan=lifespan)

# Add CORS
# In production, replace "*" with specific allowed origins
allowed_origins = os.getenv("CORS_ORIGINS", "*").split(",")
app.add_middleware(
    CORSMiddleware,
    allow_origins=allowed_origins if "*" not in allowed_origins else ["*"],
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "DELETE", "OPTIONS"],
    allow_headers=["*"],
)

if not XAI_API_KEY:
    print("Warning: XAI_API_KEY is missing. RAG features will fail.")

# Chat Client
try:
    chat_client = AsyncClient(api_key=XAI_API_KEY or "dummy_key")
except Exception:
    chat_client = None

# Management Client (needs both api_key for file uploads + management_api_key for collection ops)
try:
    if XAI_MANAGEMENT_API_KEY:
        mgmt_client = AsyncClient(
            api_key=XAI_API_KEY,
            management_api_key=XAI_MANAGEMENT_API_KEY,
        )
    else:
        print("Warning: XAI_MANAGEMENT_API_KEY missing. Collections ops will fail.")
        mgmt_client = None
except Exception:
    mgmt_client = None
    
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="token")

STATUS_PROCESSED = collections_pb2.DocumentStatus.DOCUMENT_STATUS_PROCESSED
STATUS_FAILED = collections_pb2.DocumentStatus.DOCUMENT_STATUS_FAILED

def _status_is_processed(status: object) -> bool:
    return status in ("DOCUMENT_STATUS_PROCESSED", STATUS_PROCESSED)

def _status_is_failed(status: object) -> bool:
    return status in ("DOCUMENT_STATUS_FAILED", STATUS_FAILED)

async def get_current_user(token: str = Depends(oauth2_scheme), session: AsyncSession = Depends(get_session)):
    from jose import JWTError, jwt
    from auth_utils import SECRET_KEY, ALGORITHM
    
    credentials_exception = HTTPException(
        status_code=status.HTTP_401_UNAUTHORIZED,
        detail="Could not validate credentials",
        headers={"WWW-Authenticate": "Bearer"},
    )
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        email: str = payload.get("sub")
        if email is None:
            raise credentials_exception
    except JWTError:
        raise credentials_exception
        
    statement = select(User).where(User.email == email)
    result = await session.exec(statement)
    user = result.first()
    if user is None:
        raise credentials_exception
    return user


class Filters(BaseModel):
    category: str | None = None
    tags: list[str] | None = None
    version: str | None = None
    date_from: str | None = None
    date_to: str | None = None


class ChatRequest(BaseModel):
    query: str = Field(..., min_length=1)
    collection_id: int = Field(..., description="검색할 컬렉션 ID (필수)")
    filters: Filters | None = None


class ChatResponse(BaseModel):
    request_id: str
    answer: str
    citations: list[dict] = []
    cached: bool
    latency_ms: int

class CollectionCreate(BaseModel):
    name: str
    description: str | None = None
    category: str | None = None
    tags: str | None = None  # comma-separated

class CollectionUpdate(BaseModel):
    name: str | None = None
    description: str | None = None
    category: str | None = None
    tags: str | None = None

class CollectionRead(BaseModel):
    id: int
    name: str
    xai_id: str
    description: str | None = None
    category: str | None = None
    tags: str | None = None
    created_at: str
    documents_count: int | None = None
    processing_count: int | None = None
    failed_count: int | None = None
    status: str | None = None

class Token(BaseModel):
    access_token: str
    token_type: str

class UserCreate(BaseModel):
    email: str
    password: str
    full_name: Optional[str] = None

class UserRead(BaseModel):
    id: int
    email: str
    full_name: Optional[str] = None

@app.post("/token", response_model=Token)
async def login_for_access_token(form_data: OAuth2PasswordRequestForm = Depends(), session: AsyncSession = Depends(get_session)):
    statement = select(User).where(User.email == form_data.username)
    result = await session.exec(statement)
    user = result.first()
    
    if not user or not verify_password(form_data.password, user.hashed_password):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Incorrect username or password",
            headers={"WWW-Authenticate": "Bearer"},
        )
    
    access_token_expires = timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    access_token = create_access_token(
        data={"sub": user.email}, expires_delta=access_token_expires
    )
    return {"access_token": access_token, "token_type": "bearer"}

@app.post("/register", response_model=UserRead)
async def register(user_in: UserCreate, session: AsyncSession = Depends(get_session)):
    statement = select(User).where(User.email == user_in.email)
    result = await session.exec(statement)
    if result.first():
        raise HTTPException(status_code=400, detail="Email already registered")
        
    user = User(
        email=user_in.email,
        hashed_password=get_password_hash(user_in.password),
        full_name=user_in.full_name
    )
    session.add(user)
    await session.commit()
    await session.refresh(user)
    return UserRead(id=user.id, email=user.email, full_name=user.full_name)

@app.get("/users/me", response_model=UserRead)
async def read_users_me(current_user: User = Depends(get_current_user)):
    return UserRead(id=current_user.id, email=current_user.email, full_name=current_user.full_name)


@app.get("/health")
async def health():
    return {"ok": True, "model": XAI_MODEL}


def _extract_text(content: bytes, filename: str) -> str:
    """Extract text from file content for AI analysis."""
    ext = os.path.splitext(filename.lower())[1]
    if ext in ('.txt', '.md'):
        return content.decode('utf-8', errors='replace')
    if ext in ('.docx', '.doc'):
        try:
            import docx, io
            doc = docx.Document(io.BytesIO(content))
            return "\n".join(para.text for para in doc.paragraphs if para.text.strip())
        except Exception:
            return content.decode('utf-8', errors='replace')
    if ext == '.pdf':
        try:
            import PyPDF2, io
            reader = PyPDF2.PdfReader(io.BytesIO(content))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception:
            return f"[PDF 파일: {filename}]"
    if ext in ('.doc', '.docx'):
        return f"[문서 파일: {filename}]"
    if ext in ('.jpg', '.jpeg', '.png', '.gif'):
        return f"[이미지 파일: {filename}]"
    return f"[파일: {filename}]"


ANALYZE_SYSTEM_PROMPT = """당신은 문서 온톨로지 구축을 돕는 전문 AI 어시스턴트입니다.
사용자가 업로드한 문서의 내용을 분석하여 온톨로지 메타데이터를 추천합니다.

반드시 아래 JSON 형식으로만 응답하세요. 다른 텍스트는 포함하지 마세요:
{
  "category": "문서 카테고리 (예: 정책, 재무, 기술, 법률, 인사 등)",
  "tags": ["태그1", "태그2", "태그3"],
  "summary": "문서 내용 2-3줄 요약",
  "consulting": "이 문서를 온톨로지에 통합할 때 고려할 점과 추천사항 (관련 카테고리, 연결 가능한 문서 유형, 활용 방안 등)"
}"""


@app.post("/analyze")
async def analyze_document(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
):
    """AI를 사용해 문서를 분석하고 온톨로지 메타데이터를 추천합니다."""
    if not XAI_API_KEY:
        raise HTTPException(status_code=500, detail="API Key가 설정되지 않았습니다.")

    content = await file.read()
    filename = file.filename or "unknown"

    text = _extract_text(content, filename)
    # Truncate to avoid token limits
    if len(text) > 8000:
        text = text[:8000] + "\n...(이하 생략)"

    user_msg = f"파일명: {filename}\n\n내용:\n{text}"

    try:
        async with httpx.AsyncClient(timeout=60) as client:
            resp = await client.post(
                "https://api.x.ai/v1/chat/completions",
                headers={
                    "Authorization": f"Bearer {XAI_API_KEY}",
                    "Content-Type": "application/json",
                },
                json={
                    "model": XAI_MODEL,
                    "messages": [
                        {"role": "system", "content": ANALYZE_SYSTEM_PROMPT},
                        {"role": "user", "content": user_msg},
                    ],
                    "temperature": 0.3,
                },
            )
            resp.raise_for_status()
            data = resp.json()
            answer = data["choices"][0]["message"]["content"]

            # Parse JSON from response (handle markdown code blocks)
            cleaned = answer.strip()
            if cleaned.startswith("```"):
                cleaned = cleaned.split("\n", 1)[1] if "\n" in cleaned else cleaned[3:]
                cleaned = cleaned.rsplit("```", 1)[0]
            result = json.loads(cleaned)

            return {
                "category": result.get("category", ""),
                "tags": result.get("tags", []),
                "summary": result.get("summary", ""),
                "consulting": result.get("consulting", ""),
            }
    except json.JSONDecodeError:
        # If JSON parsing fails, return the raw answer as consulting
        return {
            "category": "",
            "tags": [],
            "summary": "",
            "consulting": answer if 'answer' in dir() else "분석 결과를 파싱하지 못했습니다.",
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI 분석 실패: {str(e)}")

COLLECTION_ANALYZE_PROMPT = """당신은 문서 컬렉션 온톨로지 구축을 돕는 전문 AI 어시스턴트입니다.
컬렉션에 속한 문서 목록과 컬렉션 정보를 분석하여 적합한 온톨로지 메타데이터를 추천합니다.

반드시 아래 JSON 형식으로만 응답하세요. 다른 텍스트는 포함하지 마세요:
{
  "description": "컬렉션 설명 (2-3줄로 컬렉션의 목적과 내용 요약)",
  "category": "컬렉션 카테고리 (예: 정책, 재무, 기술, 법률, 인사, 연구 등)",
  "tags": "태그1, 태그2, 태그3 (쉼표로 구분)",
  "consulting": "이 컬렉션을 온톨로지에 통합할 때 고려할 점과 추천사항"
}"""


@app.post("/collections/{collection_id}/analyze")
async def analyze_collection(
    collection_id: int,
    session: AsyncSession = Depends(get_session),
    current_user: User = Depends(get_current_user),
):
    """컬렉션의 문서 목록을 AI로 분석하여 온톨로지 메타데이터를 추천합니다."""
    if not XAI_API_KEY:
        raise HTTPException(status_code=500, detail="API Key가 설정되지 않았습니다.")

    collection = await session.get(Collection, collection_id)
    if not collection:
        raise HTTPException(status_code=404, detail="Collection not found")

    # Gather document names
    statement = select(Document).where(Document.collection_id == collection_id)
    results = await session.exec(statement)
    documents = results.all()

    doc_list = "\n".join(f"- {d.name} (상태: {d.status})" for d in documents) if documents else "(문서 없음)"

    user_msg = (
        f"컬렉션 이름: {collection.name}\n"
        f"현재 설명: {collection.description or '(없음)'}\n"
        f"현재 카테고리: {collection.category or '(없음)'}\n"
        f"현재 태그: {collection.tags or '(없음)'}\n"
        f"문서 수: {len(documents)}\n\n"
        f"포함된 문서 목록:\n{doc_list}"
    )

    try:
        async with httpx.AsyncClient(timeout=60) as client:
            resp = await client.post(
                "https://api.x.ai/v1/chat/completions",
                headers={
                    "Authorization": f"Bearer {XAI_API_KEY}",
                    "Content-Type": "application/json",
                },
                json={
                    "model": XAI_MODEL,
                    "messages": [
                        {"role": "system", "content": COLLECTION_ANALYZE_PROMPT},
                        {"role": "user", "content": user_msg},
                    ],
                    "temperature": 0.3,
                },
            )
            resp.raise_for_status()
            data = resp.json()
            answer = data["choices"][0]["message"]["content"]

            cleaned = answer.strip()
            if cleaned.startswith("```"):
                cleaned = cleaned.split("\n", 1)[1] if "\n" in cleaned else cleaned[3:]
                cleaned = cleaned.rsplit("```", 1)[0]
            result = json.loads(cleaned)

            return {
                "description": result.get("description", ""),
                "category": result.get("category", ""),
                "tags": result.get("tags", ""),
                "consulting": result.get("consulting", ""),
            }
    except json.JSONDecodeError:
        return {
            "description": "",
            "category": "",
            "tags": "",
            "consulting": answer if 'answer' in dir() else "분석 결과를 파싱하지 못했습니다.",
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI 분석 실패: {str(e)}")


@app.get("/stats")
async def get_stats(
    session: AsyncSession = Depends(get_session),
    current_user: User = Depends(get_current_user)
):
    # Total collections
    result = await session.exec(select(func.count(Collection.id)))
    total_collections = result.one()
    
    # Total documents
    result = await session.exec(select(func.count(Document.id)))
    total_documents = result.one()
    
    # Active users (just total users for now)
    result = await session.exec(select(func.count(User.id)))
    total_users = result.one()
    
    # Total queries (from usage events)
    result = await session.exec(select(func.count(UsageEvent.id)))
    total_queries = result.one()

    # Average latency (ms)
    result = await session.exec(select(func.avg(UsageEvent.latency_ms)))
    avg_latency = result.one() or 0

    # Total cost (USD)
    result = await session.exec(select(func.sum(UsageEvent.cost_usd)))
    total_cost = result.one() or 0

    return {
        "collections": total_collections,
        "documents": total_documents,
        "users": total_users,
        "queries": total_queries,
        "avg_latency_ms": int(avg_latency),
        "cost_usd": float(total_cost),
    }

@app.get("/collections", response_model=list[CollectionRead])
async def list_collections(
    session: AsyncSession = Depends(get_session),
    current_user: User = Depends(get_current_user) # Require login
):
    result = await session.exec(select(Collection))
    collections = result.all()
    # verify format
    output: list[CollectionRead] = []
    for c in collections:
        total_docs = (await session.exec(
            select(func.count(Document.id)).where(Document.collection_id == c.id)
        )).one()
        processing_docs = (await session.exec(
            select(func.count(Document.id)).where(
                (Document.collection_id == c.id) & (Document.status == "processing")
            )
        )).one()
        failed_docs = (await session.exec(
            select(func.count(Document.id)).where(
                (Document.collection_id == c.id) & (Document.status == "failed")
            )
        )).one()
        output.append(
            CollectionRead(
                id=c.id,
                name=c.name,
                xai_id=c.xai_id,
                description=c.description,
                category=c.category,
                tags=c.tags,
                created_at=c.created_at.isoformat(),
                documents_count=total_docs,
                processing_count=processing_docs,
                failed_count=failed_docs,
                status="processing" if processing_docs > 0 else "active",
            )
        )
    return output

@app.post("/collections", response_model=CollectionRead)
async def create_collection(
    collection: CollectionCreate, 
    session: AsyncSession = Depends(get_session),
    current_user: User = Depends(get_current_user)
):
    if not mgmt_client:
         raise HTTPException(status_code=500, detail="Management API Key not configured")

    # Create in xAI
    try:
        resp = await mgmt_client.collections.create(name=collection.name)
        xai_id = resp.collection_id
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"xAI Error: {str(e)}")

    # Create in DB
    db_collection = Collection(
        name=collection.name,
        xai_id=xai_id,
        description=collection.description,
        category=collection.category,
        tags=collection.tags,
    )
    session.add(db_collection)
    await session.commit()
    await session.refresh(db_collection)

    return CollectionRead(
        id=db_collection.id,
        name=db_collection.name,
        xai_id=db_collection.xai_id,
        description=db_collection.description,
        category=db_collection.category,
        tags=db_collection.tags,
        created_at=db_collection.created_at.isoformat()
    )

@app.put("/collections/{collection_id}", response_model=CollectionRead)
async def update_collection(
    collection_id: int,
    body: CollectionUpdate,
    session: AsyncSession = Depends(get_session),
    current_user: User = Depends(get_current_user),
):
    collection = await session.get(Collection, collection_id)
    if not collection:
        raise HTTPException(status_code=404, detail="Collection not found")

    if body.name is not None:
        collection.name = body.name.strip()
    if body.description is not None:
        collection.description = body.description.strip() or None
    if body.category is not None:
        collection.category = body.category.strip() or None
    if body.tags is not None:
        collection.tags = body.tags.strip() or None

    session.add(collection)
    await session.commit()
    await session.refresh(collection)

    total_docs = (await session.exec(
        select(func.count(Document.id)).where(Document.collection_id == collection.id)
    )).one()

    return CollectionRead(
        id=collection.id,
        name=collection.name,
        xai_id=collection.xai_id,
        description=collection.description,
        category=collection.category,
        tags=collection.tags,
        created_at=collection.created_at.isoformat(),
        documents_count=total_docs,
    )

@app.delete("/collections/{collection_id}")
async def delete_collection(
    collection_id: int,
    session: AsyncSession = Depends(get_session),
    current_user: User = Depends(get_current_user)
):
    if not mgmt_client:
         raise HTTPException(status_code=500, detail="Management API Key not configured")

    collection = await session.get(Collection, collection_id)
    if not collection:
        raise HTTPException(status_code=404, detail="Collection not found")
        
    # Attempt to delete from xAI
    try:
        await mgmt_client.collections.delete(collection_id=collection.xai_id)
    except Exception as e:
        print(f"Warning: Failed to delete collection from xAI: {e}")
        # Proceed with DB deletion
    
    # Use direct SQL delete for robustness
    try:
        # Delete dependent documents
        await session.exec(delete(Document).where(Document.collection_id == collection_id))
        # Delete collection
        await session.exec(delete(Collection).where(Collection.id == collection_id))
        await session.commit()
    except Exception as e:
        print(f"Error deleting collection from DB: {e}")
        raise HTTPException(status_code=500, detail=f"Database Delete Error: {e}")
    
    return {"status": "deleted", "id": collection_id}

class DocumentRead(BaseModel):
    id: int
    name: str
    xai_doc_id: str
    status: str
    created_at: str

@app.get("/collections/{collection_id}", response_model=list[DocumentRead])
async def get_collection_documents(
    collection_id: int,
    session: AsyncSession = Depends(get_session),
    refresh: bool = False,
    current_user: User = Depends(get_current_user)
):
    collection = await session.get(Collection, collection_id)
    if not collection:
        raise HTTPException(status_code=404, detail="Collection not found")
    
    # Check if documents relationship is loaded or query them
    # Since we are using SQLModel async, we likely need to query explicitly if not joined
    # Or just select docs
    statement = select(Document).where(Document.collection_id == collection_id)
    results = await session.exec(statement)
    documents = results.all()

    if refresh and mgmt_client and documents:
        updated = False
        for doc in documents:
            if doc.status == "processed":
                continue
            try:
                status_resp = await mgmt_client.collections.get_document(
                    doc.xai_doc_id,
                    collection.xai_id,
                )
                status = getattr(status_resp, "status", None)
                if _status_is_processed(status):
                    doc.status = "processed"
                    updated = True
                elif _status_is_failed(status):
                    doc.status = "failed"
                    updated = True
            except Exception as e:
                print(f"Warning: status check failed for {doc.xai_doc_id}: {e}")
        if updated:
            await session.commit()

    return [
        DocumentRead(
            id=d.id,
            name=d.name,
            xai_doc_id=d.xai_doc_id,
            status=d.status,
            created_at=d.created_at.isoformat()
        ) for d in documents
    ]

@app.delete("/documents/{document_id}")
async def delete_document(
    document_id: int,
    session: AsyncSession = Depends(get_session),
    current_user: User = Depends(get_current_user)
):
    if not mgmt_client:
         raise HTTPException(status_code=500, detail="Management API Key not configured")

    doc = await session.get(Document, document_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
        
    # Attempt to delete from xAI
    try:
        collection = await session.get(Collection, doc.collection_id)
        if not collection:
            raise Exception("Collection not found for document")
        await delete_collection_document(mgmt_client, collection.xai_id, doc.xai_doc_id)
    except Exception as e:
        print(f"Warning: Failed to delete from xAI: {e}")
        # Proceed to delete from DB anyway so user isn't stuck
        
    session.delete(doc)
    await session.commit()
    
    return {"status": "deleted", "id": document_id}

# File upload limits
MAX_FILE_SIZE = 100 * 1024 * 1024  # 100 MB
ALLOWED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx", ".doc", ".jpg", ".jpeg", ".png", ".gif"}

@app.post("/collections/{collection_id}/upload")
async def upload_document(
    collection_id: int,
    file: UploadFile = File(...),
    category: Optional[str] = Form(None),
    tags: Optional[str] = Form(None),
    version: Optional[str] = Form(None),
    date: Optional[str] = Form(None),
    relatedDocs: Optional[str] = Form(None),
    relationship_note: Optional[str] = Form(None),
    policy_note: Optional[str] = Form(None),
    session: AsyncSession = Depends(get_session),
    current_user: User = Depends(get_current_user)
):
    if not mgmt_client:
         raise HTTPException(status_code=500, detail="Management API Key not configured")

    # Get collection
    collection = await session.get(Collection, collection_id)
    if not collection:
        raise HTTPException(status_code=404, detail="Collection not found")

    # Validate filename
    if not file.filename:
        raise HTTPException(status_code=400, detail="Filename is required")
    
    # Validate file extension
    import os
    file_ext = os.path.splitext(file.filename.lower())[1]
    if file_ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400, 
            detail=f"File type not allowed. Allowed types: {', '.join(ALLOWED_EXTENSIONS)}"
        )

    # Read file content
    content = await file.read()
    
    # Validate file size
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=400,
            detail=f"File size exceeds maximum allowed size of {MAX_FILE_SIZE / (1024 * 1024):.0f} MB"
        )
    
    if len(content) == 0:
        raise HTTPException(status_code=400, detail="File is empty")
    
    metadata = build_metadata(
        category=category,
        tags=tags,
        version=version,
        date=date,
        related_docs=relatedDocs,
        relationship_note=relationship_note,
        policy_note=policy_note,
    )

    # Convert non-text formats to plain text for xAI indexing
    upload_name = file.filename
    upload_data = content
    file_ext = os.path.splitext(file.filename.lower())[1]
    if file_ext in ('.docx', '.doc', '.pdf'):
        extracted = _extract_text(content, file.filename)
        if extracted.strip():
            upload_data = extracted.encode('utf-8')
            upload_name = os.path.splitext(file.filename)[0] + '.txt'

    # Upload to xAI
    try:
        upload_resp = await mgmt_client.collections.upload_document(
            collection_id=collection.xai_id,
            name=upload_name,
            data=upload_data,
        )
        xai_doc_id = extract_document_id(upload_resp)
        if not xai_doc_id:
            raise Exception("Could not find document_id in upload response")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"xAI Upload Error: {str(e)}")
    
    # Save to DB
    doc = Document(
        name=file.filename,
        xai_doc_id=xai_doc_id,
        collection_id=collection.id,
        status="processing" # You might want to poll status in background
    )
    session.add(doc)
    await session.commit()
    await session.refresh(doc)
    
    return {"status": "uploaded", "document_id": doc.id, "xai_doc_id": xai_doc_id}

@app.post("/chat", response_model=ChatResponse)
async def chat(
    req: ChatRequest, 
    session: AsyncSession = Depends(get_session),
    current_user: User = Depends(get_current_user)
):
    # Determine Collection ID (required)
    db_collection = await session.get(Collection, req.collection_id)
    if not db_collection:
        raise HTTPException(status_code=404, detail="지정한 컬렉션을 찾을 수 없습니다.")
    target_xai_id = db_collection.xai_id

    request_id = str(uuid.uuid4())
    t0 = time.time()

    if db_collection:
        result = await session.exec(
            select(Document).where(Document.collection_id == db_collection.id)
        )
        docs = result.all()
        if not docs:
            latency_ms = int((time.time() - t0) * 1000)
            return ChatResponse(
                request_id=request_id,
                answer="업로드된 문서가 없습니다. 먼저 문서를 업로드해 주세요.",
                citations=[],
                cached=False,
                latency_ms=latency_ms,
            )
        if mgmt_client:
            updated = False
            for doc in docs:
                if doc.status == "processed":
                    continue
                try:
                    status_resp = await mgmt_client.collections.get_document(
                        doc.xai_doc_id,
                        db_collection.xai_id,
                    )
                    status = getattr(status_resp, "status", None)
                    if _status_is_processed(status):
                        doc.status = "processed"
                        updated = True
                    elif _status_is_failed(status):
                        doc.status = "failed"
                        updated = True
                except Exception as e:
                    print(f"Warning: status check failed for {doc.xai_doc_id}: {e}")
            if updated:
                await session.commit()

            if not any(d.status == "processed" for d in docs):
                latency_ms = int((time.time() - t0) * 1000)
                return ChatResponse(
                    request_id=request_id,
                    answer=(
                        "문서가 아직 인덱싱 중입니다. 잠시 후 다시 시도해 주세요.\n\n"
                        "1) 인덱싱 대기 (가장 흔함)\n"
                        "- 보통 몇 초~10분, 큰 파일은 30분까지 걸릴 수 있습니다.\n"
                        "- 5~10분 뒤 다시 검색해 주세요.\n\n"
                        "2) 문서 상태 확인 (추천)\n"
                        "- xAI 콘솔에서 해당 컬렉션 문서 상태가 processed인지 확인하세요.\n"
                        "- processing이면 기다리면 됩니다.\n"
                        "- failed면 파일을 다시 업로드해 주세요.\n\n"
                        "3) 빠른 점검\n"
                        "- 작은 텍스트(.txt) 1개로 업로드/검색이 되는지 테스트해 보세요.\n"
                        "- 된다면 원본 파일이 크거나 복잡해 처리 지연일 가능성이 큽니다."
                    ),
                    citations=[],
                    cached=False,
                    latency_ms=latency_ms,
                )

    filters_dict = req.filters.model_dump(exclude_none=True) if req.filters else None

    # Cache key now needs to include collection
    cached = cache_get(target_xai_id, XAI_MODEL, req.query, filters_dict)
    if cached:
        latency_ms = int((time.time() - t0) * 1000)
        return ChatResponse(
            request_id=request_id,
            answer=cached["answer"],
            citations=cached.get("citations", []),
            cached=True,
            latency_ms=latency_ms,
        )

    result = await run_rag(
        client=chat_client,
        collection_id=target_xai_id,
        query=req.query,
        filters=filters_dict,
    )

    # Track usage when not cached
    usage = result.get("usage") if isinstance(result, dict) else None
    prompt_tokens = usage.get("prompt_tokens") if usage else None
    completion_tokens = usage.get("completion_tokens") if usage else None
    total_tokens = usage.get("total_tokens") if usage else None
    cost = 0.0
    if prompt_tokens is not None and completion_tokens is not None:
        cost = (prompt_tokens / 1_000_000) * COST_PER_1M_INPUT + (completion_tokens / 1_000_000) * COST_PER_1M_OUTPUT

    try:
        usage_event = UsageEvent(
            endpoint="/chat",
            model=XAI_MODEL,
            collection_id=db_collection.id if db_collection else None,
            prompt_tokens=prompt_tokens,
            completion_tokens=completion_tokens,
            total_tokens=total_tokens,
            cost_usd=cost,
            latency_ms=result.get("latency_ms"),
            cached=False,
        )
        session.add(usage_event)
        await session.commit()
    except Exception as e:
        print(f"Warning: failed to record usage: {e}")

    cache_set(target_xai_id, XAI_MODEL, req.query, filters_dict, result)

    latency_ms = int((time.time() - t0) * 1000)
    return ChatResponse(
        request_id=request_id,
        answer=result["answer"],
        citations=result.get("citations", []),
        cached=False,
        latency_ms=latency_ms,
    )
